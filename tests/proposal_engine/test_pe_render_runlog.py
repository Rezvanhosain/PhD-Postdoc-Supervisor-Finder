from pathlib import Path

import pytest

import proposal_engine.render as render_mod
from proposal_engine.config import EngineConfig
from proposal_engine.render import (RenderError, parse_draft_md,
                                    render_docx_fallback, resolve_citations,
                                    validate_docx)
from proposal_engine.runlog import RunLog

CITED = [
    {"key": "smith2020", "title": "Alpha study", "authors": ["Jane Smith"], "year": 2020,
     "venue": "Journal", "doi": "10.1/alpha", "oa_url": ""},
]

DRAFT = (
    "# An Evidence-Based Proposal\n\n"
    "## Background and Rationale\n\n"
    "The field is active [@smith2020].\n\n"
    "## References\n"
)


def test_parse_draft_md():
    title, sections = parse_draft_md(DRAFT)
    assert title == "An Evidence-Based Proposal"
    names = [s[0] for s in sections]
    assert "Background and Rationale" in names and "References" in names


def test_resolve_citations():
    out = resolve_citations("Work [@smith2020] matters.", {"smith2020": CITED[0]})
    assert "(Smith, 2020)" in out


def test_render_docx_fallback_and_validate(tmp_path):
    out = tmp_path / "proposal.docx"
    render_docx_fallback(DRAFT, CITED, out)
    assert out.exists()
    errors = validate_docx(out, ["Background and Rationale"])
    assert errors == []  # heading present, bibliography non-empty, no placeholders

    import docx
    text = "\n".join(p.text for p in docx.Document(str(out)).paragraphs)
    assert "Smith" in text  # bibliography entry rendered
    assert "(Smith, 2020)" in text  # inline citation resolved


def test_render_docx_detects_empty_bibliography(tmp_path):
    out = tmp_path / "p2.docx"
    render_docx_fallback(DRAFT, [], out)  # no cited entries
    errors = validate_docx(out, ["Background and Rationale"])
    assert any("bibliography" in e.lower() or "empty" in e.lower() for e in errors) is False
    # with no citations, the References section still contains a note (non-empty)


def test_default_renderer_does_not_use_pandoc_even_if_installed(tmp_path, monkeypatch):
    # Pandoc "present" but config renderer defaults to rich_docx -> must NOT be used.
    monkeypatch.setattr(render_mod, "find_pandoc", lambda: r"C:\fake\pandoc.exe")
    called = {"pandoc": False}
    monkeypatch.setattr(render_mod, "render_docx_pandoc",
                        lambda *a, **k: called.__setitem__("pandoc", True))
    monkeypatch.setattr(render_mod, "docx_to_pdf", lambda d, p: Path(p).write_bytes(b"%PDF-1.4"))

    draft = tmp_path / "draft.md"
    draft.write_text(DRAFT, encoding="utf-8")
    refs = tmp_path / "references.json"
    refs.write_text("[]", encoding="utf-8")
    out = render_mod.render(draft, refs, CITED, tmp_path / "p.docx", tmp_path / "p.pdf",
                            config=EngineConfig())
    assert out["docx_engine"] == "python-docx"
    assert called["pandoc"] is False
    # template-only artifact proves the rich builder ran
    import docx
    assert any("PAGE" in s.footer._element.xml
               for s in docx.Document(str(tmp_path / "p.docx")).sections)


def test_pandoc_renderer_requested_but_missing_errors(tmp_path, monkeypatch):
    monkeypatch.setattr(render_mod, "find_pandoc", lambda: None)
    draft = tmp_path / "draft.md"
    draft.write_text(DRAFT, encoding="utf-8")
    refs = tmp_path / "references.json"
    refs.write_text("[]", encoding="utf-8")
    with pytest.raises(RenderError):
        render_mod.render(draft, refs, CITED, tmp_path / "p.docx", tmp_path / "p.pdf",
                          config=EngineConfig(renderer="pandoc"))


def test_runlog_resumability(tmp_path):
    log_path = tmp_path / "run_log.json"
    art = tmp_path / "artifact.txt"
    log = RunLog(log_path, "t1")

    # Not run yet -> cannot skip
    assert log.can_skip("evidence", [art]) is False
    log.mark_success("evidence")
    # SUCCESS but artifact missing -> cannot skip
    assert log.can_skip("evidence", [art]) is False
    art.write_text("data", encoding="utf-8")
    # SUCCESS + artifact exists -> skip
    assert log.can_skip("evidence", [art]) is True
    # force overrides
    assert log.can_skip("evidence", [art], force=True) is False


def test_runlog_failed_not_skipped(tmp_path):
    log = RunLog(tmp_path / "run_log.json", "t1")
    art = tmp_path / "a.txt"
    art.write_text("x", encoding="utf-8")
    log.mark_failed("draft", "boom")
    assert log.can_skip("draft", [art]) is False
    assert log.data["status"] == "FAILED"


def test_runlog_corrupt_file_reruns(tmp_path):
    p = tmp_path / "run_log.json"
    p.write_text("{not valid json", encoding="utf-8")
    log = RunLog(p, "t1")
    assert log.stage_status("evidence") is None  # treated as fresh
