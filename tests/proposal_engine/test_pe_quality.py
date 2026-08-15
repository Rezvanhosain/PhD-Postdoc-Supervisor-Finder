"""Quality-gate detection tests: the gate must FAIL on serious defects."""
import docx

from proposal_engine import docx_builder as db
from proposal_engine.quality import document_quality_check


def _base_doc(tmp_path, name="d.docx"):
    doc = docx.Document()
    db.configure_base_document(doc)  # A4 + Page X of Y footer
    return doc, tmp_path / name


def _save(doc, path):
    doc.save(str(path))
    return path


def test_flags_missing_gantt_and_ok_footer(tmp_path):
    doc, path = _base_doc(tmp_path)
    db._heading(doc, "1. Abstract")
    db._body_paragraph(doc, "A sufficiently long abstract paragraph with enough words to pass the short check easily.")
    report = document_quality_check(_save(doc, path))
    assert any("Gantt" in s or "timeline" in s for s in report["serious"])
    assert not any("footer" in s for s in report["serious"])  # footer present


def test_flags_markdown_artifacts(tmp_path):
    doc, path = _base_doc(tmp_path)
    db._heading(doc, "1. Abstract")
    p = doc.add_paragraph()  # raw markdown injected directly, bypassing sanitize
    p.add_run("This paragraph has **bold** markers that leaked into the body text here.")
    report = document_quality_check(_save(doc, path))
    assert any("Markdown" in s for s in report["serious"])


def test_flags_adjacent_citations(tmp_path):
    doc, path = _base_doc(tmp_path)
    db._heading(doc, "1. Background")
    p = doc.add_paragraph()
    p.add_run("Evidence (Smith, 2020)(Jones, 2021) supports this claim across the field.")
    report = document_quality_check(_save(doc, path))
    assert any("adjacent citation" in s for s in report["serious"])


def test_flags_duplicated_headings(tmp_path):
    doc, path = _base_doc(tmp_path)
    db._heading(doc, "1. Methodology")
    db._body_paragraph(doc, "Body one with enough words to avoid the short-section warning here.")
    db._heading(doc, "2. Methodology")
    db._body_paragraph(doc, "Body two with enough words to avoid the short-section warning here.")
    report = document_quality_check(_save(doc, path))
    assert any("duplicated" in s for s in report["serious"])


def test_flags_missing_footer(tmp_path):
    doc = docx.Document()  # NO configure_base_document -> no footer field
    db._heading(doc, "1. Abstract")
    db._body_paragraph(doc, "Some abstract text long enough to pass the short-section check without issue.")
    path = tmp_path / "nofooter.docx"
    report = document_quality_check(_save(doc, path))
    assert any("footer" in s for s in report["serious"])
