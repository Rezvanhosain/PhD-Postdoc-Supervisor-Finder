"""Tests for the template-matching DOCX builder and the quality gate.

These build a real .docx and reopen it, so they exercise the exact output that
reaches the client (bold runs, bullet lists, RQ/H labels, the Gantt table,
metadata block, and the Page X of Y footer)."""
import docx
import pytest

from proposal_engine.config import EngineConfig
from proposal_engine.quality import document_quality_check
from proposal_engine.render import render_docx_fallback

CITED = [
    {"key": "a2020", "title": "Alpha", "authors": ["Alice Alpha"], "year": 2020,
     "venue": "J1", "doi": "10.1/a", "oa_url": ""},
    {"key": "b2021", "title": "Beta", "authors": ["Bob Beta"], "year": 2021,
     "venue": "J2", "doi": "10.1/b", "oa_url": ""},
]

DRAFT = (
    "# Availability-Aware Federated Learning for Predictive Maintenance\n\n"
    "## Abstract\n\n"
    "Industrial systems rely on **edge** analytics [@a2020][@b2021].\n\n"
    "## Objectives\n\n"
    "- To identify availability patterns\n"
    "- To assess clustering benefit\n"
    "- To evaluate staleness-aware aggregation\n\n"
    "## Research Questions and Hypotheses\n\n"
    "RQ1: How does outage frequency affect accuracy?\n"
    "RQ2: Does clustering help under non-IID data?\n"
    "H1: Incremental learning lowers post-drift error.\n"
    "H2: Availability-aware clustering improves performance.\n\n"
    "## Timeline and Work Plan\n\n"
    "The work proceeds in phases.\n"
    "- Literature review (months 1-9)\n"
    "- Method development (months 10-30)\n"
    "- Evaluation and writing (months 31-48)\n\n"
    "## References\n"
)

SECTIONS = [
    {"key": "title", "name": "Title", "words": 12},
    {"key": "abstract", "name": "Abstract", "words": 120},
    {"key": "objectives", "name": "Objectives", "words": 60},
    {"key": "questions", "name": "Research Questions and Hypotheses", "words": 80},
    {"key": "timeline", "name": "Timeline and Work Plan", "words": 120},
    {"key": "references", "name": "References", "words": 0},
]


@pytest.fixture
def built(tmp_path):
    cfg = EngineConfig(model_provider="openai", project_duration="48 months",
                       proposal_type="PhD Research Proposal",
                       applicant_name="Jordan Rivera",
                       proposed_programme="PhD in Computer Science",
                       proposal_title="Availability-Aware Federated Learning for Predictive Maintenance",
                       sections=[dict(s) for s in SECTIONS])
    out = tmp_path / "proposal.docx"
    render_docx_fallback(DRAFT, CITED, out, config=cfg, meta=cfg.metadata)
    return out, cfg


def _all_text(path):
    d = docx.Document(str(path))
    return "\n".join(p.text for p in d.paragraphs), d


def test_no_markdown_markers_in_output(built):
    out, _ = built
    text, _ = _all_text(out)
    assert "**" not in text
    assert "[@" not in text  # citations were resolved


def test_bold_becomes_real_bold_run(built):
    out, _ = built
    d = docx.Document(str(out))
    bold_runs = [r.text for p in d.paragraphs for r in p.runs if r.bold and r.text.strip()]
    assert "edge" in bold_runs  # the **edge** span rendered as a bold run


def test_adjacent_citations_merged(built):
    out, _ = built
    text, _ = _all_text(out)
    assert "(Alpha, 2020; Beta, 2021)" in text
    assert "(Alpha, 2020)(Beta, 2021)" not in text


def test_objectives_render_as_bullets(built):
    out, _ = built
    d = docx.Document(str(out))
    bullets = [p.text for p in d.paragraphs if p.style and p.style.name == "List Bullet"]
    assert len(bullets) == 3
    assert all(b.startswith("To ") for b in bullets)


def test_questions_and_hypotheses_labelled(built):
    out, _ = built
    text, _ = _all_text(out)
    for label in ("RQ1:", "RQ2:", "H1:", "H2:"):
        assert label in text


def test_gantt_table_has_year_columns(built):
    out, _ = built
    d = docx.Document(str(out))
    assert d.tables, "no Gantt table"
    header = [c.text.strip() for c in d.tables[0].rows[0].cells]
    assert header[0].lower().startswith("phase")
    assert any("Year 1" in h for h in header)
    assert any("Year 4" in h for h in header)  # 48 months -> 4 columns
    assert len(d.tables[0].rows) == 1 + 3  # header + 3 parsed phases


def test_footer_page_field_present(built):
    out, _ = built
    d = docx.Document(str(out))
    assert any("PAGE" in s.footer._element.xml for s in d.sections)


def test_title_and_metadata_render(built):
    out, cfg = built
    text, _ = _all_text(out)
    assert cfg.proposal_title in text
    assert "PhD Research Proposal" in text  # top label
    assert "Jordan Rivera" in text          # metadata field
    assert "[FILL IN]" not in text and "None" not in text.split()


def test_no_duplicate_headings(built):
    out, _ = built
    from proposal_engine.quality import _is_heading, _strip_num
    d = docx.Document(str(out))
    names = [_strip_num(p.text) for p in d.paragraphs if _is_heading(p) and p.text.strip()]
    assert len(names) == len(set(names))


def test_quality_gate_passes_clean_document(built):
    out, cfg = built
    report = document_quality_check(out, expected_title=cfg.proposal_title,
                                    section_names=[s["name"] for s in SECTIONS])
    assert report["serious"] == []
