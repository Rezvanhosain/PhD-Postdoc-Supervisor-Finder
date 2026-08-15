"""Stage 1 — Intake: extract an optional client CV/profile to Markdown and
apply a quality gate.

Supported inputs: .txt, .md, .docx, .pdf. Poor extractions are flagged
``NEEDS_PROFILE_REVIEW`` and must block drafting.
"""
from __future__ import annotations

from pathlib import Path

MIN_CHARS = 200
MAX_REPLACEMENT_RATIO = 0.02  # >2% U+FFFD => garbled


class ProfileResult:
    def __init__(self, text: str, ok: bool, reasons: list[str], source: str = ""):
        self.text = text
        self.ok = ok
        self.reasons = reasons
        self.source = source


def _extract_txt(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="replace")


def _extract_docx(path: Path) -> str:
    import docx  # python-docx

    doc = docx.Document(str(path))
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            parts.append(" \t ".join(c.text for c in row.cells))
    return "\n".join(parts)


def _extract_pdf(path: Path) -> str:
    # Prefer pypdf; fall back to PyMuPDF (fitz) which is more robust.
    text = ""
    try:
        from pypdf import PdfReader

        reader = PdfReader(str(path))
        text = "\n".join((page.extract_text() or "") for page in reader.pages)
    except Exception:
        text = ""
    if len((text or "").strip()) < MIN_CHARS:
        try:
            import fitz  # PyMuPDF

            with fitz.open(str(path)) as doc:
                text = "\n".join(page.get_text() for page in doc)
        except Exception:
            pass
    return text


def extract_profile(path: str | Path) -> ProfileResult:
    p = Path(path)
    if not p.exists():
        return ProfileResult("", False, [f"profile file not found: {p}"], str(p))
    suffix = p.suffix.lower()
    try:
        if suffix in (".txt", ".md"):
            text = _extract_txt(p)
        elif suffix == ".docx":
            text = _extract_docx(p)
        elif suffix == ".pdf":
            text = _extract_pdf(p)
        else:
            return ProfileResult("", False, [f"unsupported profile format: {suffix}"], str(p))
    except Exception as e:  # extraction library failure is a real, visible problem
        return ProfileResult("", False, [f"extraction failed: {e}"], str(p))

    reasons = _quality_issues(text)
    return ProfileResult(text.strip(), not reasons, reasons, str(p))


def _quality_issues(text: str) -> list[str]:
    reasons: list[str] = []
    stripped = (text or "").strip()
    if len(stripped) < MIN_CHARS:
        reasons.append(f"extracted profile too short ({len(stripped)} < {MIN_CHARS} chars)")
    if stripped:
        repl = stripped.count("�")
        if repl / max(len(stripped), 1) > MAX_REPLACEMENT_RATIO:
            reasons.append("excessive replacement characters (garbled extraction)")
    # recognizable structure: at least one blank-line-separated paragraph or list
    paragraphs = [b for b in (text or "").split("\n\n") if b.strip()]
    if len(paragraphs) < 1 and "\n" not in stripped:
        reasons.append("no recognizable sections or paragraphs")
    return reasons


def write_profile_md(result: ProfileResult, out_path: str | Path) -> Path:
    out = Path(out_path)
    out.parent.mkdir(parents=True, exist_ok=True)
    header = "# Extracted Client Profile\n\n"
    if not result.ok:
        header += ("> **NEEDS_PROFILE_REVIEW** — extraction quality is poor: "
                   + "; ".join(result.reasons) + "\n\n")
    out.write_text(header + (result.text or "*(no profile provided)*"), encoding="utf-8")
    return out
