"""Stage 6 — Render.

DOCX: Pandoc (with --citeproc) when available; otherwise a self-contained
python-docx builder with a small citeproc-lite for inline citations and the
reference list. PDF: LibreOffice ``soffice`` headless, else docx2pdf/Word COM.

If no PDF route works the render fails visibly (RenderPDFUnavailable) rather
than pretending success.
"""
from __future__ import annotations

import re
import shutil
import subprocess
import sys
from pathlib import Path

from .validators import PLACEHOLDER_RE

PANDOC_INSTALL_HINT = (
    "Pandoc not found. Install it with:\n"
    "    winget install --source winget --exact --id JohnMacFarlane.Pandoc\n"
    "(the engine falls back to a built-in python-docx writer in the meantime)."
)


class RenderError(RuntimeError):
    pass


class RenderPDFUnavailable(RenderError):
    """No usable PDF engine (LibreOffice or Word) on this machine."""


# --------------------------------------------------------------------- detection
def find_pandoc() -> str | None:
    return shutil.which("pandoc")


def find_soffice() -> str | None:
    exe = shutil.which("soffice") or shutil.which("soffice.exe")
    if exe:
        return exe
    for c in (r"C:\Program Files\LibreOffice\program\soffice.exe",
              r"C:\Program Files (x86)\LibreOffice\program\soffice.exe"):
        if Path(c).exists():
            return c
    return None


def word_available() -> bool:
    if sys.platform != "win32":
        return False
    try:
        import docx2pdf  # noqa: F401
    except Exception:
        return False
    for c in (r"C:\Program Files\Microsoft Office\root\Office16\WINWORD.EXE",
              r"C:\Program Files (x86)\Microsoft Office\root\Office16\WINWORD.EXE"):
        if Path(c).exists():
            return True
    return shutil.which("winword") is not None


def pdf_engine() -> str | None:
    if find_soffice():
        return "libreoffice"
    if word_available():
        return "word"
    return None


# --------------------------------------------------------------------- markdown parse
def parse_draft_md(md: str) -> tuple[str, list[tuple[str, str]]]:
    """Split our generated draft into (title, [(heading, body), ...])."""
    title = ""
    sections: list[tuple[str, list[str]]] = []
    for line in md.splitlines():
        if line.startswith("# ") and not line.startswith("## "):
            title = line[2:].strip()
        elif line.startswith("## "):
            sections.append((line[3:].strip(), []))
        else:
            if sections:
                sections[-1][1].append(line)
    return title, [(h, "\n".join(b).strip()) for h, b in sections]


# --------------------------------------------------------------------- citeproc-lite
_CITE_SPAN_RE = re.compile(r"\[([^\]]*@[^\]]+)\]")
_KEY_RE = re.compile(r"@([A-Za-z0-9_:\-]+)")


def _authors_short(authors: list[str]) -> str:
    fams = [a.split()[-1] if a.split() else a for a in authors]
    if not fams:
        return "Anon."
    if len(fams) == 1:
        return fams[0]
    if len(fams) == 2:
        return f"{fams[0]} & {fams[1]}"
    return f"{fams[0]} et al."


def _inline(entry: dict) -> str:
    return f"({_authors_short(entry.get('authors', []))}, {entry.get('year', 'n.d.')})"


def resolve_citations(text: str, by_key: dict[str, dict]) -> str:
    def repl(m: re.Match) -> str:
        keys = _KEY_RE.findall(m.group(1))
        cites = [_inline(by_key[k]).strip("()") for k in keys if k in by_key]
        return "(" + "; ".join(cites) + ")" if cites else m.group(0)
    return _CITE_SPAN_RE.sub(repl, text)


def _authors_apa(authors: list[str]) -> str:
    def fmt(a: str) -> str:
        parts = a.split()
        if len(parts) < 2:
            return a
        return f"{parts[-1]}, {' '.join(p[0] + '.' for p in parts[:-1])}"
    fs = [fmt(a) for a in authors[:20]]
    if not fs:
        return "Anon."
    if len(fs) == 1:
        return fs[0]
    return ", ".join(fs[:-1]) + ", & " + fs[-1]


def bibliography_entry(entry: dict) -> str:
    a = _authors_apa(entry.get("authors", []))
    y = entry.get("year", "n.d.")
    t = entry.get("title", "")
    v = entry.get("venue", "")
    tail = ""
    if entry.get("doi"):
        tail = f" https://doi.org/{entry['doi']}"
    elif entry.get("oa_url"):
        tail = f" {entry['oa_url']}"
    venue = f" {v}." if v else ""
    return f"{a} ({y}). {t}.{venue}{tail}".strip()


# --------------------------------------------------------------------- DOCX builders
def render_docx_fallback(md: str, cited_entries: list[dict], out_path: Path,
                         reference_docx: Path | None = None, *,
                         config=None, meta: dict | None = None) -> Path:
    """Render the styled proposal DOCX with the template-matching builder.

    Injects the citeproc-lite helpers so ``docx_builder`` (which must not import
    ``render``) can resolve inline citations and format the bibliography.
    """
    from .docx_builder import build_proposal_docx

    return build_proposal_docx(
        md, cited_entries, Path(out_path),
        config=config, meta=meta, reference_docx=reference_docx,
        resolve_citations=resolve_citations,
        bibliography_entry=bibliography_entry,
        authors_short=_authors_short,
    )


def render_docx_pandoc(draft_md_path: Path, references_json: Path, out_path: Path,
                       csl: Path | None, reference_docx: Path | None) -> Path:
    pandoc = find_pandoc()
    if not pandoc:
        raise RenderError("pandoc not available")
    cmd = [pandoc, str(draft_md_path), "--citeproc",
           "--bibliography", str(references_json), "-o", str(out_path)]
    if csl and Path(csl).exists():
        cmd += ["--csl", str(csl)]
    if reference_docx and Path(reference_docx).exists():
        cmd += ["--reference-doc", str(reference_docx)]
    try:
        subprocess.run(cmd, check=True, capture_output=True, text=True, timeout=180)
    except (subprocess.CalledProcessError, subprocess.TimeoutExpired) as e:
        detail = getattr(e, "stderr", "") or str(e)
        raise RenderError(f"pandoc failed: {detail}") from e
    return out_path


# --------------------------------------------------------------------- PDF
def docx_to_pdf(docx_path: Path, pdf_path: Path) -> Path:
    engine = pdf_engine()
    if engine == "libreoffice":
        soffice = find_soffice()
        cmd = [soffice, "--headless", "--convert-to", "pdf", "--outdir",
               str(pdf_path.parent), str(docx_path)]
        try:
            subprocess.run(cmd, check=True, capture_output=True, text=True, timeout=180)
        except (subprocess.CalledProcessError, subprocess.TimeoutExpired) as e:
            raise RenderError(f"LibreOffice PDF conversion failed: {e}") from e
        produced = pdf_path.parent / (docx_path.stem + ".pdf")
        if produced != pdf_path and produced.exists():
            produced.replace(pdf_path)
        if not pdf_path.exists():
            raise RenderError("LibreOffice reported success but no PDF was produced")
        return pdf_path
    if engine == "word":
        try:
            from docx2pdf import convert

            convert(str(docx_path), str(pdf_path))
        except Exception as e:
            raise RenderError(f"Word/docx2pdf PDF conversion failed: {e}") from e
        if not pdf_path.exists():
            raise RenderError("docx2pdf reported success but no PDF was produced")
        return pdf_path
    raise RenderPDFUnavailable(
        "No PDF engine available (LibreOffice 'soffice' not found and Microsoft "
        "Word/docx2pdf unavailable). Install LibreOffice or run on a machine with "
        "Word to produce PDFs.")


# --------------------------------------------------------------------- orchestration + validation
def render(draft_md_path: Path, references_json: Path, cited_entries: list[dict],
           docx_path: Path, pdf_path: Path, csl: Path | None = None,
           reference_docx: Path | None = None, *, config=None,
           meta: dict | None = None) -> dict:
    """Render DOCX (pandoc if available, else fallback) then PDF.

    Returns {docx, pdf, docx_engine, pdf_engine, pdf_error}. Raises RenderError
    if the DOCX cannot be produced. A missing PDF engine is reported via
    RenderPDFUnavailable to the caller (not swallowed).
    """
    md = Path(draft_md_path).read_text(encoding="utf-8")
    # The rich python-docx builder is the DEFAULT and preferred renderer: it is
    # the only path that reproduces the template (metadata page, numbered blue
    # headings, Times New Roman body, teal Gantt, Page X of Y footer). Pandoc is
    # used ONLY when config explicitly asks for it — never auto-selected just
    # because it is installed.
    renderer = (getattr(config, "renderer", None) or "rich_docx")
    if renderer == "pandoc":
        if not find_pandoc():
            raise RenderError(
                "renderer 'pandoc' was requested in config but Pandoc was not "
                "found on PATH. " + PANDOC_INSTALL_HINT)
        docx_engine = "pandoc"
        render_docx_pandoc(Path(draft_md_path), Path(references_json), docx_path,
                           csl, reference_docx)
    else:
        docx_engine = "python-docx"
        render_docx_fallback(md, cited_entries, docx_path, reference_docx,
                             config=config, meta=meta)
    if not docx_path.exists():
        raise RenderError("DOCX was not produced")

    result = {"docx": str(docx_path), "pdf": None, "docx_engine": docx_engine,
              "pdf_engine": pdf_engine(), "pdf_error": None}
    docx_to_pdf(docx_path, pdf_path)  # raises RenderPDFUnavailable / RenderError
    result["pdf"] = str(pdf_path)
    return result


def validate_docx(docx_path: Path, section_names: list[str]) -> list[str]:
    """Reopen the DOCX and check headings + non-empty bibliography + no placeholders."""
    import docx

    errors: list[str] = []
    try:
        doc = docx.Document(str(docx_path))
    except Exception as e:
        return [f"DOCX does not reopen: {e}"]
    all_text = "\n".join(p.text for p in doc.paragraphs)
    for name in section_names:
        if name.lower() in ("title", "references"):
            continue
        if name.lower() not in all_text.lower():
            errors.append(f"missing section heading: {name}")
    # bibliography non-empty: find the (possibly numbered) 'References' heading
    # and ensure a following paragraph. Headings render as "N. References".
    norm = [re.sub(r"^\s*\d+(?:\.\d+)*\.?\s+", "", p.text.strip()).lower()
            for p in doc.paragraphs]
    if "references" in norm:
        idx = norm.index("references")
        following = [t for t in norm[idx + 1:] if t]
        if not following:
            errors.append("bibliography is empty")
    else:
        errors.append("no References heading found")
    if PLACEHOLDER_RE.search(all_text):
        errors.append("placeholder markers present in rendered DOCX")
    return errors


def validate_pdf(pdf_path: Path, min_pages: int = 3) -> list[str]:
    if not pdf_path.exists():
        return [f"PDF not found: {pdf_path}"]
    try:
        from pypdf import PdfReader

        pages = len(PdfReader(str(pdf_path)).pages)
    except Exception as e:
        return [f"PDF unreadable: {e}"]
    if pages <= min_pages:
        return [f"PDF has {pages} pages (need > {min_pages})"]
    return []
