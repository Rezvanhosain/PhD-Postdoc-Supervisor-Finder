"""Rich python-docx builder that renders a proposal to match the academic
template (Times New Roman body, dark-teal numbered headings, metadata first
page, teal Gantt work-plan table, and a "Page X of Y" footer).

Colours and sizes are taken directly from the reference template:
  main headings  #0f4761 (14pt)   subheadings #2e74b5 (13pt)
  Gantt header   #04365b (navy)    active bars #52b2c0 (teal)

All heading/colour styling is applied at the run level so the output is exact
regardless of the base document's named styles; a reference.docx (built by
``build_reference_docx``) still carries matching styles for the Pandoc path.
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from . import structure
from .sanitize import (bold_segments, fix_adjacent_citations, normalize_block,
                       strip_stray_markers)

# ---- palette (from the template PDF) ---------------------------------------
HEADING_MAIN = RGBColor(0x0F, 0x47, 0x61)
HEADING_SUB = RGBColor(0x2E, 0x74, 0xB5)
GANTT_HEADER = "04365B"
GANTT_BAR = "52B2C0"
BODY_FONT = "Times New Roman"
HEADING_FONT = "Segoe UI"  # professional non-Calibri heading face (per spec)
BODY_PT = 12
MAIN_PT = 14
SUB_PT = 13

# Section keys that render as structured lists rather than paragraphs.
BULLET_KEYS = {"objectives"}
RQ_KEYS = {"questions"}
TIMELINE_KEYS = {"timeline", "work_plan", "workplan"}


# --------------------------------------------------------------------- low-level
def _set_run(run, *, font=BODY_FONT, size=BODY_PT, bold=False, color=None):
    run.font.name = font
    run.font.size = Pt(size)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color
    # ensure east-asian/complex fallback also uses the face
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs"):
        rfonts.set(qn(attr), font)
    return run


def _shade(element_pr, hex_fill: str):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_fill)
    element_pr.append(shd)


def _add_field(paragraph, instr: str, *, bold=False):
    """Append a Word field (e.g. PAGE / NUMPAGES) to a paragraph."""
    run = paragraph.add_run()
    _set_run(run, bold=bold)
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr_el = OxmlElement("w:instrText")
    instr_el.set(qn("xml:space"), "preserve")
    instr_el.text = f" {instr} "
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr_el)
    run._r.append(fld_end)


def configure_base_document(doc) -> None:
    """A4, 1-inch margins, Times New Roman 12 Normal, and a Page X of Y footer."""
    normal = doc.styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = Pt(BODY_PT)
    normal.element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), BODY_FONT)
    for sec in doc.sections:
        sec.page_height = Cm(29.7)
        sec.page_width = Cm(21.0)
        sec.top_margin = sec.bottom_margin = Cm(2.54)
        sec.left_margin = sec.right_margin = Cm(2.54)
        footer = sec.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.text = ""
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        _set_run(p.add_run("Page "))
        _add_field(p, "PAGE", bold=True)
        _set_run(p.add_run(" of "))
        _add_field(p, "NUMPAGES", bold=True)


# --------------------------------------------------------------------- blocks
def _heading(doc, text: str, *, level: int = 1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    color = HEADING_MAIN if level == 1 else HEADING_SUB
    size = MAIN_PT if level == 1 else SUB_PT
    _set_run(p.add_run(text), font=HEADING_FONT, size=size, bold=True, color=color)
    return p


def _body_paragraph(doc, text: str, *, style: str | None = None, align_justify=True):
    """Add a body paragraph, converting **bold** into real bold runs."""
    p = doc.add_paragraph(style=style) if style else doc.add_paragraph()
    if style is None:
        p.paragraph_format.space_after = Pt(8)
        p.paragraph_format.line_spacing = 1.15
        if align_justify:
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for seg, is_bold in bold_segments(text) or [(text, False)]:
        seg = seg if is_bold else strip_stray_markers(seg)
        if seg:
            _set_run(p.add_run(seg), bold=is_bold)
    return p


def _render_bullets(doc, text: str):
    for item in structure.split_bullets(text):
        _body_paragraph(doc, normalize_block(item), style="List Bullet")


def _render_labeled(doc, text: str):
    """Render Research Questions (RQ1..) and, if present, Hypotheses (H1..)."""
    # If a 'Hypothes' marker exists, split the block there so both label sets render.
    lower = text.lower()
    hyp_idx = lower.find("hypoth")
    rq_text, h_text = text, ""
    if "h1:" in lower or "h1." in lower:
        # keep hypotheses that use H labels together with RQ handling below
        pass
    if hyp_idx != -1 and "\n" in text[hyp_idx - 40:hyp_idx + 1]:
        rq_text, h_text = text[:hyp_idx], text[hyp_idx:]
    for item in structure.split_labeled(rq_text, "RQ"):
        _body_paragraph(doc, normalize_block(item), align_justify=False)
    if h_text.strip():
        for item in structure.split_labeled(h_text, "H"):
            _body_paragraph(doc, normalize_block(item), align_justify=False)


def build_gantt(doc, phases, columns, *, caption="Figure 1. Proposed Timeline for completing PhD"):
    """Render the work-plan table: navy header, teal active cells, caption."""
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(6)
    # Serif navy title, matching the template's Gantt caption styling. A navy
    # (non-heading) colour also keeps the quality gate from treating this
    # centred title as a numbered section heading.
    _set_run(title.add_run("PhD Work Plan and Timeline"), font=BODY_FONT,
             size=15, bold=True, color=RGBColor(0x04, 0x36, 0x5B))

    table = doc.add_table(rows=1, cols=1 + len(columns))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    try:
        table.style = "Table Grid"
    except KeyError:  # base doc without the built-in style
        pass
    hdr = table.rows[0].cells
    _fill_header_cell(hdr[0], "Phase / Activity")
    for i, (label, _s, _e) in enumerate(columns):
        _fill_header_cell(hdr[i + 1], label)

    for phase in phases:
        row = table.add_row().cells
        cell = row[0]
        cell.paragraphs[0].text = ""
        _set_run(cell.paragraphs[0].add_run(phase[0]), size=11)
        for i, col in enumerate(columns):
            c = row[i + 1]
            c.text = ""
            if structure.phase_active_in_year(phase, col):
                _shade(c._tc.get_or_add_tcPr(), GANTT_BAR)

    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_before = Pt(4)
    r = cap.add_run(caption)
    _set_run(r, size=10, bold=False)
    r.italic = True
    return table


def _fill_header_cell(cell, text: str):
    _shade(cell._tc.get_or_add_tcPr(), GANTT_HEADER)
    p = cell.paragraphs[0]
    p.text = ""
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_run(p.add_run(text), font=HEADING_FONT, size=11, bold=True,
             color=RGBColor(0xFF, 0xFF, 0xFF))


def _metadata_page(doc, meta: dict, top_label: str, title_text: str):
    top = doc.add_paragraph()
    top.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_run(top.add_run(top_label), size=MAIN_PT, bold=True)

    fields = [
        ("Applicant name", meta.get("applicant_name")),
        ("Proposed Programme", meta.get("proposed_programme")),
        ("Target Supervisor", meta.get("target_supervisor")),
        ("University", meta.get("university")),
        ("Date", meta.get("proposal_date")),
    ]
    any_field = any(v for _, v in fields)
    if any_field:
        doc.add_paragraph()
        for label, value in fields:
            if not value:
                continue  # omit cleanly, never a placeholder
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            _set_run(p.add_run(f"{label}: "), bold=True)
            _set_run(p.add_run(str(value)))

    doc.add_paragraph()
    _heading(doc, "Proposal Title", level=1)
    tp = doc.add_paragraph()
    tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_run(tp.add_run(title_text), size=MAIN_PT, bold=True)


# --------------------------------------------------------------------- md parse
def parse_sections(md: str) -> tuple[str, list[tuple[str, str]]]:
    title = ""
    sections: list[tuple[str, list[str]]] = []
    for line in md.splitlines():
        if line.startswith("# ") and not line.startswith("## "):
            title = line[2:].strip()
        elif line.startswith("## "):
            sections.append((line[3:].strip(), []))
        elif sections:
            sections[-1][1].append(line)
    return title, [(h, "\n".join(b).strip()) for h, b in sections]


# --------------------------------------------------------------------- main entry
def build_proposal_docx(md: str, cited_entries: list[dict], out_path: Path, *,
                        config=None, meta: dict | None = None,
                        reference_docx: Path | None = None,
                        resolve_citations=None, bibliography_entry=None,
                        authors_short=None) -> Path:
    """Assemble the styled proposal DOCX. ``resolve_citations`` / ``bibliography_entry``
    / ``authors_short`` are injected from ``render`` to avoid an import cycle."""
    meta = meta or {}
    by_key = {e["key"]: e for e in cited_entries}
    title, sections = parse_sections(md)

    name_to_key = {}
    total_months = 36
    if config is not None:
        name_to_key = {s.name.lower(): s.key for s in config.sections}
        total_months = structure.parse_duration_months(getattr(config, "project_duration", "36 months"))

    doc = (Document(str(reference_docx)) if reference_docx and Path(reference_docx).exists()
           else Document())
    configure_base_document(doc)

    proposal_title = (meta.get("proposal_title") or title or "Research Proposal").strip()
    top_label = meta.get("top_label") or getattr(config, "proposal_type", None) or "Research Proposal"
    _metadata_page(doc, meta, top_label, proposal_title)

    num = 0
    for heading, body in sections:
        key = name_to_key.get(heading.lower(), "")
        num += 1
        _heading(doc, f"{num}. {heading}", level=1)

        if heading.lower() == "references" or key == "references":
            _render_bibliography(doc, cited_entries, bibliography_entry, authors_short)
            continue

        resolved = resolve_citations(body, by_key) if resolve_citations else body
        resolved = fix_adjacent_citations(resolved)

        heading_l = heading.lower()
        is_bullets = key in BULLET_KEYS or ("objective" in heading_l and key == "")
        is_rq = key in RQ_KEYS or (("question" in heading_l or "hypoth" in heading_l) and key == "")
        is_timeline = key in TIMELINE_KEYS or (("timeline" in heading_l or "work plan" in heading_l) and key == "")

        if is_bullets:
            _render_bullets(doc, resolved)
        elif is_rq:
            _render_labeled(doc, resolved)
        elif is_timeline:
            _render_timeline_section(doc, resolved, total_months)
        else:
            for block in [b for b in resolved.split("\n\n") if b.strip()]:
                stripped = block.strip()
                if structure._BULLET_LINE_RE.match(stripped.splitlines()[0]):
                    _render_bullets(doc, stripped)
                else:
                    _body_paragraph(doc, normalize_block(" ".join(stripped.split("\n"))))

    doc.save(str(out_path))
    return out_path


def _render_timeline_section(doc, text: str, total_months: int):
    columns = structure.year_columns(total_months)
    phases = structure.parse_timeline_phases(text, total_months)
    if len(phases) < 3:
        phases = structure.default_phases(total_months)
    # Build the narrative by dropping phase bullet lines (after normalising any
    # inline ' - ' separators) so phases are not printed twice — once as prose
    # and once in the Gantt table below.
    normalized = structure._split_inline_bullets(text)
    narrative = "\n".join(ln for ln in normalized.splitlines()
                          if ln.strip() and not structure._PHASE_LINE_RE.match(ln))
    para = " ".join(narrative.split())
    if para:
        _body_paragraph(doc, normalize_block(para))
    build_gantt(doc, phases, columns)


def _render_bibliography(doc, cited_entries, bibliography_entry, authors_short):
    if not cited_entries:
        _body_paragraph(doc, "(no citations were used in the proposal body)")
        return
    key = (lambda x: authors_short(x.get("authors", [])).lower()) if authors_short else (lambda x: "")
    for e in sorted(cited_entries, key=key):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.left_indent = Cm(1.0)
        p.paragraph_format.first_line_indent = Cm(-1.0)  # hanging indent
        _set_run(p.add_run(bibliography_entry(e) if bibliography_entry else str(e)))


# --------------------------------------------------------------------- reference.docx
def build_reference_docx(path: str | Path) -> Path:
    """Generate the reusable reference.docx carrying the template's styles.

    Used by the Pandoc render path (``--reference-doc``). The python-docx builder
    styles at the run level and does not depend on this file, but it is a
    required deliverable and keeps both render paths visually consistent.
    """
    from docx.enum.style import WD_STYLE_TYPE

    out = Path(path)
    out.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure_base_document(doc)

    def style_heading(name, size, color):
        try:
            st = doc.styles[name]
        except KeyError:
            st = doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        st.font.name = HEADING_FONT
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = color
        st.element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), HEADING_FONT)

    style_heading("Heading 1", MAIN_PT, HEADING_MAIN)
    style_heading("Heading 2", SUB_PT, HEADING_SUB)
    # Title style used by the metadata page top label
    try:
        tstyle = doc.styles["Title"]
        tstyle.font.name = BODY_FONT
        tstyle.font.size = Pt(MAIN_PT)
        tstyle.font.bold = True
    except KeyError:
        pass
    doc.save(str(out))
    return out
