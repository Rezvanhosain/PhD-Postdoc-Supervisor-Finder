"""Proposal quality gate.

``document_quality_check`` reopens a rendered DOCX and reports defects that would
make it look like raw AI/Markdown output rather than a finished academic
proposal. Serious issues block "client-deliverable" status; warnings are advisory.

It is deliberately conservative: it inspects the *rendered* document (paragraph
text, run colours, tables, footer fields), so it catches problems no matter how
they arose in the pipeline.
"""
from __future__ import annotations

from pathlib import Path

from .docx_builder import HEADING_MAIN, HEADING_SUB
from .sanitize import has_adjacent_citations, has_markdown_artifacts

_HEADING_RGBS = {str(HEADING_MAIN), str(HEADING_SUB)}
LONG_PARAGRAPH_WORDS = 260
SHORT_SECTION_WORDS = 12


def _is_heading(p) -> bool:
    for r in p.runs:
        try:
            rgb = r.font.color.rgb
        except Exception:
            rgb = None
        if rgb is not None and str(rgb) in _HEADING_RGBS:
            return True
    return False


def _footer_has_page_field(doc) -> bool:
    for sec in doc.sections:
        xml = sec.footer._element.xml if sec.footer is not None else ""
        if "PAGE" in xml:
            return True
    return False


def _has_gantt(doc) -> bool:
    for t in doc.tables:
        if not t.rows:
            continue
        first = t.rows[0].cells[0].text.strip().lower()
        if "phase" in first and "activ" in first:
            return True
    return False


def document_quality_check(docx_path: str | Path, *, expected_title: str | None = None,
                           section_names: list[str] | None = None) -> dict:
    """Return ``{"serious": [...], "warnings": [...]}`` for a rendered DOCX."""
    import docx

    serious: list[str] = []
    warnings: list[str] = []
    try:
        doc = docx.Document(str(docx_path))
    except Exception as e:  # unreadable file is itself a serious defect
        return {"serious": [f"DOCX does not reopen: {e}"], "warnings": []}

    paras = doc.paragraphs
    headings = [(i, p) for i, p in enumerate(paras) if _is_heading(p) and p.text.strip()]
    body_text = "\n".join(p.text for p in paras if not _is_heading(p))

    # 1) raw Markdown artifacts in body
    if has_markdown_artifacts(body_text):
        serious.append("raw Markdown markers (**, __, or #) present in body text")
    # 2) adjacent citations
    if has_adjacent_citations(body_text):
        serious.append("adjacent citation groups without separator, e.g. (A, 2016)(B, 2018)")
    # 3) duplicated headings (same normalised name appears twice, or twice in a row)
    norm = [_strip_num(p.text) for _, p in headings]
    dupes = sorted({n for n in norm if norm.count(n) > 1 and n})
    if dupes:
        serious.append(f"duplicated section headings: {dupes}")

    # 4) proposal title present
    all_text = "\n".join(p.text for p in paras)
    if expected_title and expected_title.strip() and expected_title.strip() not in all_text:
        serious.append("proposal title missing from document")
    if not any("proposal title" == _strip_num(p.text).lower() for _, p in headings) \
            and "Proposal Title" not in all_text:
        warnings.append("no 'Proposal Title' block found")

    # 5) Gantt / timeline table
    if not _has_gantt(doc):
        serious.append("work-plan Gantt/timeline table not found")

    # 6) footer page numbering
    if not _footer_has_page_field(doc):
        serious.append("footer 'Page X of Y' field not present")

    # 7) list-formatted sections + 8) empty/short/long sections
    blocks = _section_blocks(paras, headings)
    for name, block_paras in blocks.items():
        low = name.lower()
        words = sum(len(p.text.split()) for p in block_paras)
        if low in ("references", "proposal title"):
            continue
        if not block_paras or words == 0:
            serious.append(f"section '{name}' is empty")
            continue
        if "objective" in low:
            if not any(p.style and p.style.name == "List Bullet" for p in block_paras):
                serious.append(f"objectives section '{name}' not formatted as a bullet list")
        if "question" in low or "hypoth" in low:
            labelled = any(p.text.strip()[:3].upper().startswith(("RQ", "H1", "H2", "H3"))
                           for p in block_paras)
            if not labelled:
                serious.append(f"questions/hypotheses section '{name}' lacks RQ/H labels")
        if words < SHORT_SECTION_WORDS:
            warnings.append(f"section '{name}' is very short ({words} words)")
        for p in block_paras:
            if len(p.text.split()) > LONG_PARAGRAPH_WORDS:
                warnings.append(f"very long paragraph in '{name}' ({len(p.text.split())} words)")
                break

    return {"serious": serious, "warnings": warnings}


def _strip_num(text: str) -> str:
    import re
    return re.sub(r"^\s*\d+(?:\.\d+)*\.?\s+", "", text or "").strip()


def _section_blocks(paras, headings) -> dict:
    """Map each heading name -> the list of body paragraphs beneath it."""
    blocks: dict[str, list] = {}
    idxs = [i for i, _ in headings]
    for h_pos, (i, p) in enumerate(headings):
        name = _strip_num(p.text)
        start = i + 1
        end = idxs[h_pos + 1] if h_pos + 1 < len(idxs) else len(paras)
        blocks[name] = [paras[j] for j in range(start, end) if paras[j].text.strip()]
    return blocks
