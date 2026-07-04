"""Shared DOCX/PDF writers. Content is a list of (kind, text) blocks where
kind is 'h1' | 'h2' | 'p' | 'bullet'."""
from __future__ import annotations

from pathlib import Path

Block = tuple[str, str]


def write_docx(blocks: list[Block], path: Path) -> Path:
    from docx import Document
    from docx.shared import Pt

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    for kind, text in blocks:
        if kind == "h1":
            doc.add_heading(text, level=1)
        elif kind == "h2":
            doc.add_heading(text, level=2)
        elif kind == "bullet":
            doc.add_paragraph(text, style="List Bullet")
        else:
            doc.add_paragraph(text)
    doc.save(str(path))
    return path


def write_pdf(blocks: list[Block], path: Path) -> Path:
    from fpdf import FPDF

    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=18)
    pdf.add_page()
    for kind, text in blocks:
        safe = text.encode("latin-1", "replace").decode("latin-1")
        if kind == "h1":
            pdf.set_font("Helvetica", "B", 16)
            pdf.multi_cell(0, 9, safe, new_x="LMARGIN", new_y="NEXT")
            pdf.ln(2)
        elif kind == "h2":
            pdf.set_font("Helvetica", "B", 13)
            pdf.multi_cell(0, 8, safe, new_x="LMARGIN", new_y="NEXT")
            pdf.ln(1)
        elif kind == "bullet":
            pdf.set_font("Helvetica", "", 11)
            pdf.multi_cell(0, 6, "  - " + safe, new_x="LMARGIN", new_y="NEXT")
        else:
            pdf.set_font("Helvetica", "", 11)
            pdf.multi_cell(0, 6, safe, new_x="LMARGIN", new_y="NEXT")
            pdf.ln(1)
    pdf.output(str(path))
    return path
